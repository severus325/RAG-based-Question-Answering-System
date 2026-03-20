from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


@dataclass(slots=True)
class RawDocument:
    source: str
    text: str


def _normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def read_pdf_file(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def read_docx_file(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def load_documents(directory: Path) -> tuple[list[RawDocument], list[str]]:
    documents: list[RawDocument] = []
    notes: list[str] = []

    for path in sorted(directory.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            if path.suffix.lower() in {".txt", ".md"}:
                text = read_text_file(path)
            elif path.suffix.lower() == ".pdf":
                text = read_pdf_file(path)
            else:
                text = read_docx_file(path)
        except ModuleNotFoundError as exc:
            notes.append(f"{path.name}: 缺少依赖 {exc.name}，该文件暂未载入。")
            continue
        except Exception as exc:  # noqa: BLE001
            notes.append(f"{path.name}: 读取失败，原因: {exc}")
            continue

        normalized = _normalize_text(text)
        if not normalized:
            notes.append(f"{path.name}: 文档内容为空，已跳过。")
            continue

        documents.append(RawDocument(source=str(path.relative_to(directory)), text=normalized))

    if not documents:
        notes.append("文档目录中还没有可检索内容。")

    return documents, notes

